# encoding:utf-8

from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


document = Document()

document.add_heading("This is the tile", 0)
p = document.add_paragraph("The Gravity Recovery and Climate Experiment (GRACE) satellites, joint lunched by the National Aeronautics and Space Administration (NASA) in the United States and the Deutsche Forschungsanstalt für Luft-und Raumfahrt (DLR), measure monthly territorial water storage anomalies (TWSA) by converting observed gravity anomalies into changes of equivalent water height (EWH) (Rodell et al. 1999, Wiese et al. 2016, Zhao et al. 2019). There are two different approaches to derive the TWSA from GRACE data: the spherical harmonic coefficients (SHC) solutions based on GRACE Level-2 gravity field, and the mass concentration functions (mascon) solutions based on GRACE Level 1B data(Swenson et al. 2006, Luthcke et al. 2013, Watkins et al. 2015). Compared with SHC solutions, mascon solutions offer several key advantages: 1) mascon solutions are easily implemented to filter out noise from GRACE observations. 2) mascon solutions better handle the mass leakage problem than SHC solutions, which would be beneficial for studying the regional mass change(Richey et al. 2015, Scanlon et al. 2016). (Wang et al. 2020) pointed out that groundwater storage anomalies (GWSA) in the NCP derived from GRACE mascon solution are reliable and robust for groundwater drought identification and evaluation. So, in this study, the GRACE RL06 mascon products from the Center for Space Research (CSR) of the University of Texas were used to quantify TWSA (available on http://www2.csr.utexas.edu/grace), and then to estimate the GWSA based on water balance equation (2) (Zhang et al. 2019). The monthly CSR mascon block, with a spatial resolution of 0.25° × 0.25°, covers the period from April 2002 to December 2021, with 34 months of missing data due to technical issues. (Zhang et al. 2021) compared four different classical interpolations, and found that the piecewise interpolation performs best over the entire testing period. Therefore, the piecewise linear interpolation method was adopted to fill the missing data (Lin et al. 2019, Wu et al. 2019, Wang et al. 2020).")
p.add_run("bold").bold = True
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote') #添加项目列表（前面一个小圆点）
document.add_paragraph( 'first item in unordered list', style='List Bullet' )
document.add_paragraph('second item in unordered list', style='List Bullet') #添加项目列表（前面数字）
document.add_paragraph('first item in ordered list', style='List Number')
document.add_paragraph('second item in ordered list', style='List Number') #添加图片
document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
) #添加表格：一行三列 # 表格样式参数可选：# Normal Table # Table Grid # Light Shading、 Light Shading Accent 1 至 Light Shading Accent 6 # Light List、Light List Accent 1 至 Light List Accent 6 # Light Grid、Light Grid Accent 1 至 Light Grid Accent 6 # 太多了其它省略...
table = document.add_table(rows=1, cols=3, style='Light Shading Accent 2') #获取第一行的单元格列表
hdr_cells = table.rows[0].cells #下面三行设置上面第一行的三个单元格的文本值
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records: #表格添加行，并返回行所在的单元格列表
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break() #保存.docx文档
document.save('demo.docx')



